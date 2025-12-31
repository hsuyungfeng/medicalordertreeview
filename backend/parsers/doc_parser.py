"""
文檔解析器 - 支援 DOC/DOCX/XLSX/CSV 格式
"""

import subprocess
import hashlib
import logging
import re
from pathlib import Path
from typing import Optional, Dict, List
from datetime import datetime
from docx import Document as DocxDocument
from openpyxl import load_workbook
import pandas as pd

from models import DocumentContent, Section, TableData

logger = logging.getLogger(__name__)


class DocParser:
    """文檔解析器，支援 DOC/DOCX/XLSX 格式"""

    def __init__(self, cache_dir: Path):
        """
        初始化解析器

        Args:
            cache_dir: 緩存目錄路徑
        """
        self.cache_dir = Path(cache_dir)
        self.converted_dir = self.cache_dir / "converted"
        self.converted_dir.mkdir(parents=True, exist_ok=True)

    def parse_file(self, file_path: Path) -> Optional[DocumentContent]:
        """
        解析文件（自動偵測格式）

        Args:
            file_path: 文件路徑

        Returns:
            解析後的文檔內容，或在解析失敗時返回 None
        """
        file_path = Path(file_path)

        if not file_path.exists():
            logger.error(f"文件不存在: {file_path}")
            return None

        suffix = file_path.suffix.lower()

        try:
            if suffix == ".doc":
                # 轉換 .doc 為 .docx
                docx_path = self._convert_doc_to_docx(file_path)
                if not docx_path:
                    return None
                return self._parse_docx(docx_path, file_path)
            elif suffix == ".docx":
                return self._parse_docx(file_path, file_path)
            elif suffix == ".xlsx":
                return self._parse_xlsx(file_path)
            elif suffix == ".csv":
                return self._parse_csv(file_path)
            else:
                logger.warning(f"不支援的文件格式: {suffix}")
                return None
        except Exception as e:
            logger.error(f"解析 {file_path} 失敗: {e}", exc_info=True)
            return None

    def _convert_doc_to_docx(self, doc_path: Path) -> Optional[Path]:
        """
        使用 LibreOffice 轉換 .doc 為 .docx

        Args:
            doc_path: .doc 文件路徑

        Returns:
            轉換後的 .docx 文件路徑，失敗時返回 None
        """
        output_path = self.converted_dir / f"{doc_path.stem}.docx"

        # 檢查緩存
        if output_path.exists():
            logger.debug(f"使用緩存的轉換文件: {output_path}")
            return output_path

        try:
            logger.info(f"轉換 .doc 文件: {doc_path} -> {output_path}")
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", str(self.converted_dir),
                    str(doc_path)
                ],
                check=True,
                capture_output=True,
                timeout=30
            )

            if output_path.exists():
                logger.info(f"轉換成功: {output_path}")
                return output_path
            else:
                logger.error(f"轉換後文件不存在: {output_path}")
                return None

        except subprocess.TimeoutExpired:
            logger.error(f"轉換超時: {doc_path}")
            return None
        except FileNotFoundError:
            logger.error("未找到 soffice 命令，請安裝 LibreOffice")
            return None
        except Exception as e:
            logger.error(f"轉換失敗: {e}")
            return None

    def _parse_docx(self, docx_path: Path, original_path: Path) -> Optional[DocumentContent]:
        """
        解析 DOCX 文件

        Args:
            docx_path: DOCX 文件路徑
            original_path: 原始文件路徑（用於生成 doc_id）

        Returns:
            解析後的文檔內容
        """
        try:
            doc = DocxDocument(docx_path)
            sections = []
            char_position = 0

            # 解析段落和標題
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                # 檢查是否為標題
                style_name = para.style.name
                level = self._get_heading_level(style_name)

                section = Section(
                    id=f"sec_{len(sections)}",
                    heading=para.text if level > 0 else "",
                    level=level if level > 0 else 0,
                    content=para.text,
                    tables=[],
                    position=char_position
                )
                sections.append(section)
                char_position += len(para.text) + 1  # +1 for newline

            # 提取表格
            for table in doc.tables:
                table_data = self._parse_table(table)
                # 附加到最近的章節
                if sections:
                    sections[-1].tables.append(table_data)

            # 計算元數據
            word_count = sum(len(p.text.split()) for p in doc.paragraphs)
            page_count = len(doc.sections)

            # 計算文件哈希
            file_hash = self._compute_hash(original_path)

            # 生成 doc_id（從文件名提取）
            doc_id = original_path.stem

            return DocumentContent(
                doc_id=doc_id,
                title=doc.core_properties.title or doc_id,
                filename=original_path.name,
                sections=sections,
                metadata={
                    "page_count": page_count,
                    "word_count": word_count,
                    "table_count": len(doc.tables),
                    "section_count": len(sections)
                },
                parsed_at=datetime.now(),
                file_hash=file_hash
            )

        except Exception as e:
            logger.error(f"解析 DOCX 失敗 {docx_path}: {e}", exc_info=True)
            return None

    def _parse_xlsx(self, xlsx_path: Path) -> Optional[DocumentContent]:
        """
        解析 XLSX 文件

        Args:
            xlsx_path: XLSX 文件路徑

        Returns:
            解析後的文檔內容
        """
        try:
            wb = load_workbook(xlsx_path)
            sections = []
            char_position = 0

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]

                # 添加工作表名作為標題
                section = Section(
                    id=f"sec_{len(sections)}",
                    heading=sheet_name,
                    level=1,
                    content=f"工作表: {sheet_name}",
                    tables=[],
                    position=char_position
                )
                sections.append(section)
                char_position += len(sheet_name) + 1

                # 提取表格數據
                if ws.max_row > 0 and ws.max_column > 0:
                    headers = []
                    rows = []

                    for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                        row_data = [str(cell) if cell is not None else "" for cell in row]
                        if row_idx == 0:
                            headers = row_data
                        else:
                            rows.append(row_data)

                    if headers:
                        table_data = TableData(headers=headers, rows=rows)
                        section.tables.append(table_data)

            # 計算元數據
            file_hash = self._compute_hash(xlsx_path)
            doc_id = xlsx_path.stem

            return DocumentContent(
                doc_id=doc_id,
                title=doc_id,
                filename=xlsx_path.name,
                sections=sections,
                metadata={
                    "sheet_count": len(wb.sheetnames),
                    "section_count": len(sections)
                },
                parsed_at=datetime.now(),
                file_hash=file_hash
            )

        except Exception as e:
            logger.error(f"解析 XLSX 失敗 {xlsx_path}: {e}", exc_info=True)
            return None

    def _parse_table(self, table) -> TableData:
        """
        解析 DOCX 表格

        Args:
            table: python-docx 表格對象

        Returns:
            解析後的表格數據
        """
        headers = []
        rows = []

        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text for cell in row.cells]
            if row_idx == 0:
                headers = row_data
            else:
                rows.append(row_data)

        return TableData(headers=headers, rows=rows)

    def _parse_csv(self, csv_path: Path) -> Optional[DocumentContent]:
        """
        解析醫療給付項目 CSV

        Args:
            csv_path: CSV 文件路徑

        Returns:
            解析後的文檔內容
        """
        try:
            logger.info(f"開始解析 CSV: {csv_path}")
            df = pd.read_csv(csv_path, encoding='utf-8')
            logger.info(f"CSV 包含 {len(df)} 行")

            sections = []
            hierarchy = self._build_hierarchy_from_names(df)

            logger.info(f"建立了 {len(hierarchy)} 個分類層級")

            # 創建分類節點和項目子節點
            position = 0
            for category_idx, (category, items) in enumerate(hierarchy.items()):
                # 創建主分類節點
                main_section = Section(
                    id=f"csv_cat_{category_idx}",
                    heading=category,
                    level=1,
                    content=f"共 {len(items)} 個項目",
                    tables=[],
                    position=position
                )
                sections.append(main_section)
                position += len(category) + 1

                # 創建子項目節點
                for item_idx, (idx, row) in enumerate(items.iterrows()):
                    code = str(row['診療項目代碼']).strip()
                    name = str(row['中文項目名稱']).strip()
                    item_heading = f"{code} - {name}"

                    sub_section = Section(
                        id=f"csv_item_{code}",
                        heading=item_heading,
                        level=2,
                        content=self._format_csv_content(row),
                        tables=[self._create_payment_table(row)],
                        position=position
                    )
                    sections.append(sub_section)
                    position += len(item_heading) + 1

            # 計算元數據
            file_hash = self._compute_hash(csv_path)
            doc_id = csv_path.stem

            logger.info(f"CSV 解析完成: {len(sections)} 個章節")

            return DocumentContent(
                doc_id=doc_id,
                title="醫療服務給付項目 (CSV)",
                filename=csv_path.name,
                sections=sections,
                metadata={
                    "row_count": len(df),
                    "source_type": "csv",
                    "categories": len(hierarchy),
                    "section_count": len(sections)
                },
                parsed_at=datetime.now(),
                file_hash=file_hash
            )

        except Exception as e:
            logger.error(f"解析 CSV 失敗 {csv_path}: {e}", exc_info=True)
            return None

    def _build_hierarchy_from_names(self, df: pd.DataFrame) -> Dict[str, pd.DataFrame]:
        """
        基於中文項目名稱建立階層結構

        Args:
            df: 包含醫療項目數據的 DataFrame

        Returns:
            分類 → 項目 DataFrame 的字典
        """
        hierarchy = {}

        for idx, row in df.iterrows():
            name = str(row['中文項目名稱']).strip()
            category = self._extract_category(name)

            if category not in hierarchy:
                hierarchy[category] = []

            hierarchy[category].append(row)

        # 轉換為 DataFrame
        return {cat: pd.DataFrame(items) for cat, items in hierarchy.items()}

    def _extract_category(self, name: str) -> str:
        """
        從項目名稱提取分類（多層級支援）

        Args:
            name: 中文項目名稱

        Returns:
            分類名稱
        """
        # 按優先級排序的關鍵字分類
        # 優先級高的關鍵字應該排在前面，以避免被後續關鍵字覆蓋
        keywords = [
            ('內視鏡', '內視鏡類'),
            ('超音波', '超音波類'),
            ('放射', '放射線類'),
            ('病理檢驗', '病理檢驗類'),
            ('檢查', '檢查類'),
            ('檢驗', '檢驗類'),
            ('檢體', '檢體檢查類'),
            ('手術', '手術類'),
            ('麻醉', '麻醉類'),
            ('復健', '復健類'),
            ('處置', '處置類'),
            ('治療', '治療類'),
            ('注射', '注射類'),
            ('訪視', '訪視類'),
            ('診察', '診察類'),
            ('藥事', '藥事服務類'),
        ]

        # 按優先級查找關鍵字
        for keyword, category in keywords:
            if keyword in name:
                # 可以進一步細化多層級
                if keyword in ['檢查', '檢驗']:
                    # 進一步區分血液、尿液等檢查類型
                    if '血液' in name or '血' in name:
                        return '檢查類 > 血液檢查'
                    elif '尿液' in name or '尿' in name:
                        return '檢查類 > 尿液檢查'
                    elif '影像' in name or '放射' in name:
                        return '檢查類 > 影像檢查'
                    else:
                        return category
                return category

        # 根據代碼前綴分類（如果名稱包含代碼）
        if '-' in name:
            code_part = name.split('-')[0].strip()
            # 嘗試提取數字代碼
            code_match = re.search(r'\d+', code_part)
            if code_match:
                code = code_match.group()
                if code.startswith('0'):
                    return '基本診療類'
                elif code.startswith('1'):
                    return '門診診療類'
                elif code.startswith('2'):
                    return '住院診療類'
                elif code.startswith('3'):
                    return '牙科診療類'
                elif code.startswith('4'):
                    return '中醫診療類'
                elif code.startswith('5'):
                    return '其他服務類'

        # 檢查是否包含特殊標記
        if '已廢止' in name or '廢止' in name:
            return '已廢止項目'

        return '其他類'

    def _format_csv_content(self, row: pd.Series) -> str:
        """
        格式化 CSV 行為內容文本

        Args:
            row: CSV 數據行

        Returns:
            格式化的內容文本
        """
        code = str(row['診療項目代碼']).strip()
        points = str(row['健保支付點數']).strip()
        start_date = str(row['生效起日']).strip()
        end_date = str(row['生效迄日']).strip()

        content = f"""診療項目代碼：{code}
健保支付點數：{points} 點
生效日期：{start_date} ~ {end_date}"""

        # 添加英文名稱（如果有）
        if pd.notna(row['英文項目名稱']) and str(row['英文項目名稱']).strip():
            english_name = str(row['英文項目名稱']).strip()
            if english_name and english_name.lower() != 'nan':
                content += f"\n英文名稱：{english_name}"

        # 添加支付規定（如果有）
        if pd.notna(row['支付規定']):
            rules = str(row['支付規定']).strip()
            if rules and rules.lower() != 'nan':
                # 格式化支付規定文本
                formatted_rules = self._format_payment_rules(rules)
                content += f"\n\n支付規定：\n{formatted_rules}"

        return content

    def _format_payment_rules(self, rules_text: str) -> str:
        """
        格式化支付規定文本，提取和組織審查原則

        Args:
            rules_text: 原始支付規定文本

        Returns:
            格式化後的規定文本
        """
        if not rules_text or rules_text.lower() == 'nan':
            return ""

        # 提取審查原則段落
        lines = rules_text.split('\n')
        formatted_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 移除項目代碼前綴（如 "(檢查)06002C - ..."）
            if line.startswith('('):
                close_paren = line.find(')')
                if close_paren > 0:
                    line = line[close_paren + 1:].strip()
                    # 移除開頭的 "-"
                    if line.startswith('-'):
                        line = line[1:].strip()

            # 檢查是否為審查原則標題
            if '審查原則' in line or '申報條件' in line or '應檢附' in line or '限制' in line:
                if formatted_lines:
                    formatted_lines.append("")  # 空行分隔
                formatted_lines.append(f"【{line}】")
            else:
                # 編號項目（如 "1.", "2.", 等）
                if re.match(r'^\d+\.', line):
                    formatted_lines.append(f"  {line}")
                else:
                    if formatted_lines and not formatted_lines[-1].startswith('【'):
                        formatted_lines.append(f"  {line}")
                    else:
                        formatted_lines.append(line)

        return '\n'.join(formatted_lines) if formatted_lines else rules_text

    def _create_payment_table(self, row: pd.Series) -> TableData:
        """
        創建支付信息表格

        Args:
            row: CSV 數據行

        Returns:
            TableData 對象
        """
        code = str(row['診療項目代碼']).strip()
        points = str(row['健保支付點數']).strip()
        start_date = str(row['生效起日']).strip()
        end_date = str(row['生效迄日']).strip()

        return TableData(
            headers=['項目', '內容'],
            rows=[
                ['診療項目代碼', code],
                ['支付點數', f"{points} 點"],
                ['生效起日', start_date],
                ['生效迄日', end_date]
            ]
        )

    def _get_heading_level(self, style_name: str) -> int:
        """
        從樣式名提取標題層級

        Args:
            style_name: 樣式名稱（如 'Heading 1', 'Heading 2'）

        Returns:
            層級（1-6），非標題返回 0
        """
        if style_name.startswith("Heading "):
            try:
                return int(style_name.split()[-1])
            except (ValueError, IndexError):
                return 0
        return 0

    def _compute_hash(self, file_path: Path) -> str:
        """
        計算文件 SHA256 哈希值

        Args:
            file_path: 文件路徑

        Returns:
            SHA256 哈希值（十六進制字符串）
        """
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()
